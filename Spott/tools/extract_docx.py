#!/usr/bin/env python3
"""Extract a DOCX into readable Markdown while preserving paragraph/table order."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError(f"Unsupported parent: {type(parent)!r}")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def clean(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split())


def escape_cell(text: str) -> str:
    return clean(text).replace("|", "\\|")


def paragraph_markdown(paragraph: Paragraph) -> str:
    # This source document wraps most text runs in structured document tags
    # (w:sdt), which python-docx's Paragraph.text does not currently expose.
    text = clean("".join(node.text or "" for node in paragraph._p.xpath(".//w:t")))
    if not text:
        return ""
    style = paragraph.style.name if paragraph.style else ""
    lowered = style.lower()
    if lowered.startswith("title"):
        return f"# {text}"
    if lowered.startswith("subtitle"):
        return f"_{text}_"
    if lowered.startswith("heading"):
        digits = "".join(ch for ch in style if ch.isdigit())
        level = max(1, min(6, int(digits or "1")))
        return f"{'#' * level} {text}"
    return f"[{style}] {text}" if style and style != "Normal" else text


def table_markdown(table: Table, index: int) -> list[str]:
    rows = [
        [
            escape_cell(
                " ".join(
                    clean("".join(node.text or "" for node in p.xpath(".//w:t")))
                    for p in cell._tc.xpath(".//w:p")
                )
            )
            for cell in row.cells
        ]
        for row in table.rows
    ]
    if not rows:
        return [f"<!-- TABLE {index}: empty -->"]
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    output = [f"<!-- TABLE {index}: {len(rows)} rows x {width} cols -->"]
    output.append("| " + " | ".join(rows[0]) + " |")
    output.append("| " + " | ".join(["---"] * width) + " |")
    output.extend("| " + " | ".join(row) + " |" for row in rows[1:])
    return output


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()

    doc = Document(args.input)
    output: list[str] = [f"<!-- source: {args.input} -->", ""]
    table_index = 0
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            line = paragraph_markdown(block)
            if line:
                output.extend([line, ""])
        else:
            table_index += 1
            output.extend(table_markdown(block, table_index))
            output.append("")

    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text("\n".join(output), encoding="utf-8")

    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
    print(f"output={args.output} chars={args.output.stat().st_size}")


if __name__ == "__main__":
    main()
