#!/usr/bin/env python3
"""Build the polished Spott development specification DOCX from Markdown."""

from __future__ import annotations

import argparse
import math
import re
import tempfile
import zipfile
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree


# compact_reference_guide preset + named Spott brand overrides.
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN = {"top": 80, "bottom": 80, "start": 120, "end": 120}

INK = "17181C"
MUTED = "6F737C"
PURPLE = "6E5BE7"
PURPLE_DARK = "3C326F"
PURPLE_PALE = "EFECFF"
CORAL = "FF745F"
MINT = "3DBD91"
AMBER = "D99A2B"
DANGER = "D84B5B"
CANVAS = "F7F5F0"
SURFACE = "FFFFFF"
DIVIDER = "DAD7D0"
TABLE_HEADER = "E8EEF5"  # exact compact_reference_guide fill
TABLE_ALT = "FAFAFC"
CODE_BG = "F4F3F8"
CODE_BORDER = "D8D4E7"
CALLOUT_BG = "F4F1FF"

LATIN_FONT = "Noto Sans SC"  # embedded cross-platform font; covers Latin, Chinese and Japanese
CJK_FONT = "Noto Sans SC"
MONO_FONT = "Menlo"
DIAGRAM_FONT_PATH = "/System/Library/Fonts/Hiragino Sans GB.ttc"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
FONT_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, side: str, color: str, size: int = 8, space: int = 6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    border = borders.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        borders.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size: float | None = None, color: str | None = None,
                 bold: bool | None = None, italic: bool | None = None,
                 name: str = LATIN_FONT, east_asia: str = CJK_FONT) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size: float, color: str = INK, bold: bool = False,
                   name: str = LATIN_FONT, east_asia: str = CJK_FONT) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)
    instr_run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    instr_run.append(instr)
    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)
    text_run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = placeholder
    text_run.append(text)
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph._p.append(begin_run)
    paragraph._p.append(instr_run)
    paragraph._p.append(separate_run)
    paragraph._p.append(text_run)
    paragraph._p.append(end_run)


def add_hyperlink(paragraph, text: str, url: str, color: str = PURPLE) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), CJK_FONT)
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|https?://[^\s）)]+)")


def add_inline(paragraph, text: str, *, default_size: float = 11, default_color: str = INK,
               bold: bool = False) -> None:
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, default_size, default_color, bold=bold)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, default_size - 0.4, PURPLE_DARK, bold=False, name=MONO_FONT, east_asia=CJK_FONT)
            set_paragraph_shading(paragraph, "F8F7FB")
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, default_size, default_color, bold=True)
        else:
            add_hyperlink(paragraph, token, token)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, default_size, default_color, bold=bold)


def set_table_geometry(table, widths: list[int], indent: int = TABLE_INDENT_DXA) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in CELL_MARGIN.items():
                el = tc_mar.find(qn(f"w:{side}"))
                if el is None:
                    el = OxmlElement(f"w:{side}")
                    tc_mar.append(el)
                el.set(qn("w:w"), str(value))
                el.set(qn("w:type"), "dxa")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "D9DCE3")


def compute_widths(rows: list[list[str]]) -> list[int]:
    col_count = max(len(r) for r in rows)
    stats = []
    for col in range(col_count):
        values = [r[col] if col < len(r) else "" for r in rows]
        measured = [sum(2 if ord(ch) > 127 else 1 for ch in v) for v in values]
        max_len = min(54, max([8] + measured))
        head_len = measured[0] if measured else 8
        stats.append(max(5.0, (max_len ** 0.70) + head_len * 0.05))
    minimum = 1050 if col_count <= 4 else 800
    available = CONTENT_WIDTH_DXA - minimum * col_count
    total_weight = sum(stats)
    widths = [minimum + int(available * weight / total_weight) for weight in stats]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def set_table_cell_text(cell, text: str, *, header: bool, font_size: float) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.style = "Normal"
    p.paragraph_format.space_before = Pt(1 if not header else 2)
    p.paragraph_format.space_after = Pt(1 if not header else 2)
    p.paragraph_format.line_spacing = 1.15
    add_inline(p, text.replace("\\|", "|"), default_size=font_size,
               default_color=INK if not header else PURPLE_DARK, bold=header)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    rows = [r + [""] * (col_count - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=col_count)
    widths = compute_widths(rows)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    font_size = 9.2 if col_count <= 2 else (8.7 if col_count == 3 else 8.25)
    for r_idx, row in enumerate(table.rows):
        if r_idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, TABLE_HEADER)
        elif r_idx % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, TABLE_ALT)
        for c_idx, cell in enumerate(row.cells):
            set_table_cell_text(cell, rows[r_idx][c_idx], header=r_idx == 0, font_size=font_size)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)
    spacer.paragraph_format.line_spacing = 0.5


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    raw = []
    idx = start
    while idx < len(lines) and lines[idx].lstrip().startswith("|"):
        raw.append(lines[idx].strip())
        idx += 1
    rows = []
    for line_index, line in enumerate(raw):
        cells = [c.strip() for c in re.split(r"(?<!\\)\|", line.strip("|"))]
        if line_index == 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows, idx


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.05
    set_paragraph_shading(p, CODE_BG)
    set_paragraph_border(p, "left", PURPLE, size=14, space=8)
    lines = code.rstrip().splitlines() or [""]
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, 8.3, PURPLE_DARK, name=MONO_FONT, east_asia=CJK_FONT)
        if idx < len(lines) - 1:
            run.add_break()


def add_callout(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.10)
    p.paragraph_format.right_indent = Inches(0.05)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    set_paragraph_shading(p, CALLOUT_BG)
    set_paragraph_border(p, "left", PURPLE, size=16, space=8)
    add_inline(p, text, default_size=10.2, default_color=PURPLE_DARK)


def create_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    next_abs = max(abstract_ids, default=-1) + 1

    def add_abstract(abs_id: int, fmt: str, lvl_text: str, font: str | None = None) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        txt = OxmlElement("w:lvlText")
        txt.set(qn("w:val"), lvl_text)
        lvl.append(txt)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        if font:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), font)
            fonts.set(qn("w:hAnsi"), font)
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    add_abstract(next_abs, "bullet", "•", LATIN_FONT)
    add_abstract(next_abs + 1, "decimal", "%1.")

    def add_num(abs_id: int) -> int:
        num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
        num_id = max(num_ids, default=0) + 1
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(abs_id))
        num.append(abs_ref)
        numbering.append(num)
        return num_id

    return add_num(next_abs), next_abs + 1


def add_list_item(doc: Document, text: str, num_id: int, *, font_size: float = 10.5,
                  space_after: float = 3, line_spacing: float = 1.20) -> None:
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    add_inline(p, text, default_size=font_size)


def add_picture_with_alt(doc: Document, path: Path, caption: str, alt_text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(6.25))
    doc_pr = inline._inline.docPr
    doc_pr.set("name", caption)
    doc_pr.set("descr", alt_text)
    cap = doc.add_paragraph()
    cap.style = "Caption"
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, 8.5, MUTED, italic=True)


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, 10.5, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.20

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, 16, PURPLE_DARK, True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.05
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, 13, PURPLE, True)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.line_spacing = 1.10
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, 12, PURPLE_DARK, True)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    set_style_font(caption, 8.5, MUTED, False)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.keep_with_next = True


def setup_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)


def setup_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    ht = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(ht, [5200, 4160], indent=0)
    remove_table_borders(ht)
    left, right = ht.rows[0].cells
    lp = left.paragraphs[0]
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lr = lp.add_run("SPOTT · PRODUCT ENGINEERING")
    rr = rp.add_run("V1.0 · 2026-07-15")
    set_run_font(lr, 8, PURPLE_DARK, bold=True)
    set_run_font(rr, 8, MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    ft = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(ft, [6000, 3360], indent=0)
    remove_table_borders(ft)
    left, right = ft.rows[0].cells
    lp = left.paragraphs[0]
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lr = lp.add_run("DEVELOPMENT BASELINE · INTERNAL")
    set_run_font(lr, 7.8, MUTED)
    r = rp.add_run("PAGE ")
    set_run_font(r, 7.8, MUTED)
    add_field(rp, "PAGE", "1")
    r2 = rp.add_run(" / ")
    set_run_font(r2, 7.8, MUTED)
    add_field(rp, "NUMPAGES", "—")


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("SPOTT  /  PRODUCT ENGINEERING")
    set_run_font(r, 10, PURPLE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Spott iOS 原生 App\n+ Web 全栈开发文档")
    set_run_font(r, 30, INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("面向 10 万用户的高可用、可同步、可持续演进研发基线")
    set_run_font(r, 14, MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("SWIFT 6.3  ·  NEXT.JS 16.2  ·  NESTJS 11  ·  POSTGRESQL 18.4")
    set_run_font(r, 9.5, PURPLE_DARK, bold=True)
    set_paragraph_shading(p, PURPLE_PALE)
    set_paragraph_border(p, "left", PURPLE, size=18, space=8)

    rows = [
        ["产品", "Spott - 日本同城活动"],
        ["目标", "发现活动，遇见同好"],
        ["交付范围", "iOS 原生 / Web-PWA / Backend / Ops / PostgreSQL"],
        ["数据架构", "PostgreSQL 唯一权威源 · Web 与 iOS 双向同步"],
        ["版本", "V1.0 · 2026-07-15"],
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [1900, 7460])
    for idx, row in enumerate(table.rows):
        set_cell_shading(row.cells[0], CANVAS)
        set_table_cell_text(row.cells[0], rows[idx][0], header=True, font_size=9)
        set_table_cell_text(row.cells[1], rows[idx][1], header=False, font_size=9.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("DESIGN DIRECTION")
    set_run_font(r, 8.5, CORAL, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Tokyo Afterglow")
    set_run_font(r, 21, PURPLE_DARK, bold=True)
    p = doc.add_paragraph()
    r = p.add_run("借鉴 Luma 的高级感与内容优先原则，建立 Spott 独立的城市活动设计语言。")
    set_run_font(r, 10, MUTED)

    doc.add_page_break()


def add_toc(doc: Document, main_lines: list[str]) -> None:
    p = doc.add_paragraph("目录", style="Heading 1")
    p.paragraph_format.page_break_before = True
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(12)
    add_inline(intro, "共 20 个主章节与 4 个附录；详细层级可在 Word 导航窗格中查看。", default_size=9.5, default_color=MUTED)
    entries = []
    for line in main_lines:
        match = re.match(r"^#\s+(.+)$", line.strip())
        if match:
            title = match.group(1)
            number_match = re.match(r"^(\d+\.|附录\s+[A-D]\.)\s*(.*)$", title)
            if number_match:
                entries.append([number_match.group(1), number_match.group(2)])
            else:
                entries.append(["", title])
    half = math.ceil(len(entries) / 2)
    rows = [["章节", "主题", "章节", "主题"]]
    for index in range(half):
        left = entries[index]
        right = entries[index + half] if index + half < len(entries) else ["", ""]
        rows.append(left + right)
    add_markdown_table(doc, rows)
    doc.add_page_break()


def add_heading(doc: Document, level: int, text: str, *, top_level_page_break: bool = True) -> None:
    p = doc.add_paragraph(text, style=f"Heading {min(level, 3)}")
    if level == 1 and top_level_page_break:
        p.paragraph_format.page_break_before = True
        set_paragraph_border(p, "bottom", PURPLE_PALE, size=10, space=5)


def parse_markdown(doc: Document, lines: list[str], diagrams: dict[str, tuple[Path, str, str]],
                   bullet_num_id: int, decimal_abs_id: int, *, top_break: bool = True) -> None:
    numbering = doc.part.numbering_part.element

    def new_decimal_num() -> int:
        num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
        num_id = max(num_ids, default=0) + 1
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), str(decimal_abs_id))
        num.append(ref)
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        override.append(start_override)
        num.append(override)
        numbering.append(num)
        return num_id

    idx = 0
    in_code = False
    code_lines: list[str] = []
    current_decimal_id: int | None = None
    previous_was_decimal = False
    compact_reference_list = False
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, "\n".join(code_lines))
                in_code = False
                code_lines = []
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue
        if not stripped:
            previous_was_decimal = False
            idx += 1
            continue

        diagram_match = re.fullmatch(r"\[\[DIAGRAM:([^]]+)\]\]", stripped)
        if diagram_match:
            key = diagram_match.group(1)
            path, caption, alt = diagrams[key]
            add_picture_with_alt(doc, path, caption, alt)
            idx += 1
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            title = heading.group(2)
            if level == 1:
                compact_reference_list = title.startswith("附录 C.")
            add_heading(doc, level, title, top_level_page_break=top_break)
            idx += 1
            continue

        if stripped.startswith("|") and idx + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-{3,}", lines[idx + 1]):
            rows, idx = parse_table(lines, idx)
            add_markdown_table(doc, rows)
            continue

        if stripped.startswith(">"):
            quote = []
            while idx < len(lines) and lines[idx].strip().startswith(">"):
                quote.append(lines[idx].strip()[1:].strip())
                idx += 1
            add_callout(doc, "  ·  ".join(quote))
            continue

        if stripped.startswith("- "):
            add_list_item(
                doc,
                stripped[2:],
                bullet_num_id,
                font_size=8.7 if compact_reference_list else 10.5,
                space_after=1.5 if compact_reference_list else 3,
                line_spacing=1.08 if compact_reference_list else 1.20,
            )
            previous_was_decimal = False
            idx += 1
            continue

        ordered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if ordered:
            if not previous_was_decimal or current_decimal_id is None:
                current_decimal_id = new_decimal_num()
            add_list_item(doc, ordered.group(2), current_decimal_id)
            previous_was_decimal = True
            idx += 1
            continue

        previous_was_decimal = False
        p = doc.add_paragraph()
        p.style = "Normal"
        add_inline(p, stripped, default_size=10.5)
        idx += 1


def font(size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(DIAGRAM_FONT_PATH, size=size)


def draw_centered(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str,
                  font_obj: ImageFont.FreeTypeFont, fill: str, spacing: int = 8) -> None:
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 40
    lines: list[str] = []
    for raw in text.split("\n"):
        current = ""
        for ch in raw:
            candidate = current + ch
            if draw.textlength(candidate, font=font_obj) > max_width and current:
                lines.append(current)
                current = ch
            else:
                current = candidate
        lines.append(current)
    bbox = draw.multiline_textbbox((0, 0), "\n".join(lines), font=font_obj, spacing=spacing, align="center")
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.multiline_text(((x1 + x2 - width) / 2, (y1 + y2 - height) / 2), "\n".join(lines),
                        font=font_obj, fill=fill, spacing=spacing, align="center")


def rounded_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str,
                fill: str, outline: str = DIVIDER, text_fill: str = INK,
                radius: int = 24, font_size: int = 28) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=f"#{fill}", outline=f"#{outline}", width=3)
    draw_centered(draw, box, text, font(font_size), f"#{text_fill}")


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = PURPLE) -> None:
    draw.line([start, end], fill=f"#{color}", width=5)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    for delta in (2.55, -2.55):
        p = (end[0] + length * math.cos(angle + delta), end[1] + length * math.sin(angle + delta))
        draw.line([end, p], fill=f"#{color}", width=5)


def diagram_canvas(title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (1800, 1000), f"#{CANVAS}")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), title, font=font(44), fill=f"#{PURPLE_DARK}")
    draw.line((70, 110, 1730, 110), fill=f"#{PURPLE}", width=5)
    return image, draw


def make_architecture(path: Path) -> None:
    image, draw = diagram_canvas("Spott 高可用总体架构")
    clients = [(90, 160, 520, 260), (685, 160, 1115, 260), (1280, 160, 1710, 260)]
    for box, label in zip(clients, ("iOS Native\nSwift 6.3", "Web / PWA\nNext.js 16.2", "Ops Web\nMFA + RBAC")):
        rounded_box(draw, box, label, SURFACE, PURPLE_PALE, PURPLE_DARK, font_size=28)
    rounded_box(draw, (210, 325, 1590, 420), "Route 53  →  CloudFront / WAF  →  ALB（三可用区）", PURPLE_PALE, PURPLE, PURPLE_DARK, font_size=30)
    apps = [(100, 500, 520, 620), (690, 500, 1110, 620), (1280, 500, 1700, 620)]
    for box, label in zip(apps, ("API\nNestJS + Fastify", "Realtime Gateway\nWebSocket / SSE", "Async Workers\nOutbox Consumers")):
        rounded_box(draw, box, label, SURFACE, DIVIDER, INK, font_size=27)
    data = [(70, 735, 480, 895), (510, 735, 890, 895), (920, 735, 1300, 895), (1330, 735, 1730, 895)]
    labels = ("PostgreSQL 18.4\n唯一权威事实源\n1 Writer + 2 Standby", "Valkey / Redis\nMulti-AZ Cache", "SQS + DLQ\n至少一次投递", "S3 + CDN\n媒体与证据")
    fills = (PURPLE_PALE, SURFACE, SURFACE, SURFACE)
    outlines = (PURPLE, DIVIDER, DIVIDER, DIVIDER)
    for box, label, fill_value, outline in zip(data, labels, fills, outlines):
        rounded_box(draw, box, label, fill_value, outline, PURPLE_DARK if fill_value == PURPLE_PALE else INK, font_size=25)
    for x in (305, 900, 1495):
        arrow(draw, (x, 260), (x, 325))
    for x in (305, 900, 1495):
        arrow(draw, (x, 420), (x, 500))
    for x in (305, 900, 1495):
        arrow(draw, (x, 620), (x, 735))
    image.save(path, quality=95)


def make_sync(path: Path) -> None:
    image, draw = diagram_canvas("Web 与 iOS 可靠增量同步")
    rounded_box(draw, (80, 180, 500, 350), "iOS 本地库\nSwiftData + Offline Queue", SURFACE, PURPLE, PURPLE_DARK, font_size=28)
    rounded_box(draw, (80, 610, 500, 780), "Web Cache\nQuery Cache + IndexedDB Draft", SURFACE, PURPLE, PURPLE_DARK, font_size=27)
    rounded_box(draw, (680, 250, 1120, 710), "Versioned API\n\nIdempotency-Key\nbaseVersion\ndeviceId\n\nPush / Pull Cursor", PURPLE_PALE, PURPLE, PURPLE_DARK, font_size=29)
    rounded_box(draw, (1300, 180, 1720, 410), "PostgreSQL Transaction\n实体 + version\nchange_log + outbox", SURFACE, PURPLE, INK, font_size=27)
    rounded_box(draw, (1300, 570, 1720, 800), "Realtime Hint\nWebSocket / APNs\n只唤醒，不充当事实日志", SURFACE, CORAL, INK, font_size=27)
    arrow(draw, (500, 265), (680, 360))
    arrow(draw, (500, 695), (680, 600))
    arrow(draw, (1120, 380), (1300, 300))
    arrow(draw, (1300, 690), (1120, 610))
    arrow(draw, (1510, 410), (1510, 570), CORAL)
    draw.text((690, 820), "关键链路强一致 · 展示数据最终一致 · 断线后按 cursor 追赶", font=font(27), fill=f"#{MUTED}")
    image.save(path, quality=95)


def make_data_domains(path: Path) -> None:
    image, draw = diagram_canvas("PostgreSQL 领域 Schema 与权威数据边界")
    rounded_box(draw, (600, 360, 1200, 650), "PostgreSQL 18.4\n\nConstraints · Transactions\nRLS · Audit · PITR", PURPLE_PALE, PURPLE, PURPLE_DARK, font_size=31)
    boxes = [
        ((70, 170, 430, 300), "identity\n账号 / 会话 / 资料"),
        ((500, 150, 850, 290), "events\n活动 / 报名 / 签到"),
        ((950, 150, 1300, 290), "community\n群组 / 内容 / 成就"),
        ((1370, 170, 1730, 300), "commerce\n钱包 / 账本 / 订单"),
        ((70, 730, 430, 860), "safety\n举报 / 证据 / 审核"),
        ((500, 750, 850, 890), "notification\n通知 / 投递 / 偏好"),
        ((950, 750, 1300, 890), "admin + growth\n配置 / 审计 / 归因"),
        ((1370, 730, 1730, 860), "sync\nchange_log / outbox"),
    ]
    centers = [(300, 300), (675, 290), (1125, 290), (1500, 300), (300, 730), (675, 750), (1125, 750), (1500, 730)]
    targets = [(600, 430), (740, 360), (1060, 360), (1200, 430), (600, 580), (740, 650), (1060, 650), (1200, 580)]
    for (box, label), start, end in zip(boxes, centers, targets):
        rounded_box(draw, box, label, SURFACE, DIVIDER, INK, font_size=25)
        arrow(draw, start, end, PURPLE)
    image.save(path, quality=95)


def make_event_state(path: Path) -> None:
    image, draw = diagram_canvas("活动状态机（服务端权威）")
    states = [
        ((70, 190, 320, 300), "草稿"),
        ((430, 190, 700, 300), "待审核"),
        ((820, 190, 1090, 300), "已发布"),
        ((1210, 190, 1510, 300), "报名截止"),
        ((1430, 430, 1710, 540), "进行中"),
        ((1080, 700, 1360, 810), "已结束"),
        ((570, 700, 850, 810), "已取消"),
        ((100, 700, 380, 810), "已下架"),
        ((430, 420, 700, 530), "需修改"),
    ]
    for box, label in states:
        fill_value = PURPLE_PALE if label in ("待审核", "已发布") else SURFACE
        outline = PURPLE if fill_value == PURPLE_PALE else DIVIDER
        rounded_box(draw, box, label, fill_value, outline, PURPLE_DARK if fill_value == PURPLE_PALE else INK, font_size=30)
    connections = [
        ((320, 245), (430, 245)), ((700, 245), (820, 245)), ((1090, 245), (1210, 245)),
        ((1510, 270), (1570, 430)), ((1430, 540), (1280, 700)),
        ((560, 300), (560, 420)), ((430, 475), (320, 300)),
        ((900, 300), (740, 700)), ((900, 300), (240, 700)),
    ]
    for start, end in connections:
        arrow(draw, start, end, CORAL if end[1] == 700 else PURPLE)
    draw.text((90, 900), "任何取消/下架均停止报名与分享；状态迁移同时写审计、change_log 与 outbox。", font=font(27), fill=f"#{MUTED}")
    image.save(path, quality=95)


def make_registration_state(path: Path) -> None:
    image, draw = diagram_canvas("报名、候补与签到状态机")
    boxes = {
        "填写中": (70, 180, 320, 290),
        "待确认": (430, 150, 700, 260),
        "已报名": (820, 250, 1100, 370),
        "候补中": (430, 520, 700, 630),
        "递补待确认": (820, 520, 1150, 630),
        "已签到": (1300, 180, 1580, 300),
        "未到场": (1300, 430, 1580, 550),
        "已取消/拒绝": (850, 760, 1210, 870),
        "活动取消": (250, 760, 570, 870),
    }
    for label, box in boxes.items():
        emphasis = label in ("已报名", "已签到")
        rounded_box(draw, box, label, PURPLE_PALE if emphasis else SURFACE,
                    PURPLE if emphasis else DIVIDER, PURPLE_DARK if emphasis else INK, font_size=29)
    links = [
        ("填写中", "待确认"), ("填写中", "已报名"), ("填写中", "候补中"),
        ("待确认", "已报名"), ("候补中", "递补待确认"), ("递补待确认", "已报名"),
        ("已报名", "已签到"), ("已报名", "未到场"),
        ("待确认", "已取消/拒绝"), ("候补中", "已取消/拒绝"),
        ("已报名", "活动取消"),
    ]
    for src, dst in links:
        sb = boxes[src]; db = boxes[dst]
        start = ((sb[0] + sb[2]) // 2, (sb[1] + sb[3]) // 2)
        end = ((db[0] + db[2]) // 2, (db[1] + db[3]) // 2)
        arrow(draw, start, end, CORAL if "取消" in dst or "拒绝" in dst else PURPLE)
    draw.text((90, 920), "容量、积分和状态在同一 PostgreSQL 事务内裁决；客户端不直接 PATCH status。", font=font(27), fill=f"#{MUTED}")
    image.save(path, quality=95)


def make_diagrams(output_dir: Path) -> dict[str, tuple[Path, str, str]]:
    output_dir.mkdir(parents=True, exist_ok=True)
    specs = {
        "architecture": (make_architecture, "图 1 · Spott 高可用总体架构", "iOS、Web 和运营后台通过边缘层访问 API、实时网关和异步任务，PostgreSQL 是唯一权威源。"),
        "sync_flow": (make_sync, "图 2 · Web 与 iOS 可靠增量同步", "客户端离线队列通过幂等 API 写 PostgreSQL，变更日志和实时提示驱动增量拉取。"),
        "data_domains": (make_data_domains, "图 3 · PostgreSQL 领域数据边界", "PostgreSQL 按身份、活动、社群、积分、安全、通知、运营和同步领域分区。"),
        "event_state": (make_event_state, "图 4 · 活动状态机", "活动从草稿、审核、发布到截止、进行和结束，并可进入需修改、取消或下架。"),
        "registration_state": (make_registration_state, "图 5 · 报名与签到状态机", "报名在待确认、已报名、候补、递补、签到、未到场和取消状态间受控迁移。"),
    }
    result = {}
    for key, (maker, caption, alt) in specs.items():
        path = output_dir / f"{key}.png"
        maker(path)
        result[key] = (path, caption, alt)
    return result


def embed_truetype_font(docx_path: Path, font_path: Path, font_name: str) -> None:
    """Embed a redistributable TTF so headless Office and user devices retain CJK glyphs."""
    if not font_path.is_file():
        raise FileNotFoundError(f"embedded font not found: {font_path}")

    font_table_path = "word/fontTable.xml"
    font_rels_path = "word/_rels/fontTable.xml.rels"
    settings_path = "word/settings.xml"
    content_types_path = "[Content_Types].xml"
    package_font_path = f"word/fonts/{font_path.name}"

    with zipfile.ZipFile(docx_path, "r") as source_zip:
        source_items = {item.filename: (item, source_zip.read(item.filename)) for item in source_zip.infolist()}

    font_table = etree.fromstring(source_items[font_table_path][1])
    for existing in font_table.findall(f"{{{W_NS}}}font"):
        if existing.get(f"{{{W_NS}}}name") == font_name:
            font_table.remove(existing)

    if font_rels_path in source_items:
        font_rels = etree.fromstring(source_items[font_rels_path][1])
    else:
        font_rels = etree.Element(f"{{{PKG_REL_NS}}}Relationships", nsmap={None: PKG_REL_NS})
    used_ids = {rel.get("Id") for rel in font_rels}
    rel_number = 1
    while f"rId{rel_number}" in used_ids:
        rel_number += 1
    font_rel_id = f"rId{rel_number}"
    relationship = etree.SubElement(font_rels, f"{{{PKG_REL_NS}}}Relationship")
    relationship.set("Id", font_rel_id)
    relationship.set("Type", FONT_REL_TYPE)
    relationship.set("Target", f"fonts/{font_path.name}")

    font = etree.SubElement(font_table, f"{{{W_NS}}}font")
    font.set(f"{{{W_NS}}}name", font_name)
    charset = etree.SubElement(font, f"{{{W_NS}}}charset")
    charset.set(f"{{{W_NS}}}val", "86")
    family = etree.SubElement(font, f"{{{W_NS}}}family")
    family.set(f"{{{W_NS}}}val", "swiss")
    pitch = etree.SubElement(font, f"{{{W_NS}}}pitch")
    pitch.set(f"{{{W_NS}}}val", "variable")
    for element_name in ("embedRegular", "embedBold", "embedItalic", "embedBoldItalic"):
        embedded = etree.SubElement(font, f"{{{W_NS}}}{element_name}")
        embedded.set(f"{{{W_NS}}}fontKey", "{00000000-0000-0000-0000-000000000000}")
        embedded.set(f"{{{R_NS}}}id", font_rel_id)
        embedded.set(f"{{{W_NS}}}subsetted", "0")

    settings = etree.fromstring(source_items[settings_path][1])
    for setting_name in ("embedTrueTypeFonts", "saveSubsetFonts"):
        if settings.find(f"{{{W_NS}}}{setting_name}") is None:
            settings.append(etree.Element(f"{{{W_NS}}}{setting_name}"))

    content_types = etree.fromstring(source_items[content_types_path][1])
    if not any(
        child.tag == f"{{{CONTENT_TYPES_NS}}}Default" and child.get("Extension", "").lower() == "ttf"
        for child in content_types
    ):
        default_type = etree.SubElement(content_types, f"{{{CONTENT_TYPES_NS}}}Default")
        default_type.set("Extension", "ttf")
        default_type.set("ContentType", "application/x-font-ttf")

    replacements = {
        font_table_path: etree.tostring(font_table, xml_declaration=True, encoding="UTF-8", standalone=True),
        font_rels_path: etree.tostring(font_rels, xml_declaration=True, encoding="UTF-8", standalone=True),
        settings_path: etree.tostring(settings, xml_declaration=True, encoding="UTF-8", standalone=True),
        content_types_path: etree.tostring(content_types, xml_declaration=True, encoding="UTF-8", standalone=True),
        package_font_path: font_path.read_bytes(),
    }

    with tempfile.NamedTemporaryFile(
        prefix=f".{docx_path.stem}-", suffix=".docx", dir=docx_path.parent, delete=False
    ) as temp_file:
        temp_path = Path(temp_file.name)
    try:
        with zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6) as target_zip:
            for name, (item, payload) in source_items.items():
                if name not in replacements:
                    target_zip.writestr(item, payload)
            for name, payload in replacements.items():
                target_zip.writestr(name, payload)
        temp_path.replace(docx_path)
    finally:
        temp_path.unlink(missing_ok=True)


def build(source: Path, output: Path, assets_dir: Path, embedded_font: Path) -> None:
    text = source.read_text(encoding="utf-8")
    lines = text.splitlines()
    main_index = next(i for i, line in enumerate(lines) if line.startswith("# 1. "))
    front_lines = lines[:main_index]
    main_lines = lines[main_index:]

    doc = Document()
    setup_page(doc)
    setup_styles(doc)
    setup_header_footer(doc)
    bullet_num_id, decimal_abs_id = create_numbering(doc)
    diagrams = make_diagrams(assets_dir)

    doc.core_properties.title = "Spott iOS 原生 App + Web 全栈开发文档"
    doc.core_properties.subject = "Swift 原生、Web、PostgreSQL、跨端同步与高可用研发规格"
    doc.core_properties.author = "Spott Product Engineering"
    doc.core_properties.keywords = "Spott, Swift, iOS, Web, PostgreSQL, Sync, High Availability"
    doc.core_properties.comments = "Development baseline generated from the approved product requirements."

    add_cover(doc)
    add_heading(doc, 1, "文档控制", top_level_page_break=False)
    front_filtered = []
    for line in front_lines:
        if line.startswith("# Spott ") or line.startswith("> 版本：") or line.startswith("> 适用端：") or line.startswith("> 权威数据源：") or line.startswith("> 产品口号："):
            continue
        front_filtered.append(line)
    parse_markdown(doc, front_filtered, diagrams, bullet_num_id, decimal_abs_id, top_break=False)
    add_toc(doc, main_lines)
    parse_markdown(doc, main_lines, diagrams, bullet_num_id, decimal_abs_id, top_break=True)

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    embed_truetype_font(output, embedded_font, CJK_FONT)
    print(f"saved={output}")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--assets-dir", type=Path, required=True)
    parser.add_argument("--embedded-font", type=Path, required=True)
    args = parser.parse_args()
    build(args.source, args.output, args.assets_dir, args.embedded_font)


if __name__ == "__main__":
    main()
